# -*- coding: utf-8 -*-
"""Gera RESUMO_ENGENHARIA_RESERVATORIOS_I.docx a partir do HTML fonte."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r"RESUMO_ENGENHARIA_RESERVATORIOS_I.docx"

doc = Document()

# ─── PAGE SETUP ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin   = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── HELPERS ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(text, level=1, color=None, bold=True):
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
    p = doc.add_paragraph(style=style_map.get(level, 'Heading 1'))
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def add_box(text, bg="F0F0F0", border_color="2980B9"):
    """Add a shaded formula/concept box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    return p

def add_numbered_item(text, number, bold_prefix=False):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(10.5)
    return p

def styled_table(headers, rows, header_bg="1A5276", header_fg="FFFFFF"):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(
            int(header_fg[0:2],16), int(header_fg[2:4],16), int(header_fg[4:6],16))
        set_cell_bg(hdr_cells[i], header_bg)

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg = "FFFFFF" if r_idx % 2 == 0 else "EBF5FB"
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            set_cell_bg(cells[c_idx], bg)

    return table

def section_divider(title, color_hex, text_color_tuple):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("  " + title + "  ")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*text_color_tuple)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("ENGENHARIA DE RESERVATÓRIOS I")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("RESUMO COMPLETO DE ESTUDO")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
for line in ["ISPTEC — Instituto Superior Politécnico de Tecnologias e Ciências",
             "Prof. Dr. Geraldo A. R. Ramos",
             "3º Ano · 2º Semestre · 2025 / 2026"]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph("Aulas 0 – 12  |  Capítulos 1 – 5")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Índice", 1, (0x1A, 0x1A, 0x2E))
toc_items = [
    ("Capítulo 1", "Sistema Petrolífero, Tipos de Reservatórios e Mecanismos de Produção"),
    ("Capítulo 2", "Propriedades dos Fluidos de Reservatório"),
    ("Capítulo 3", "Propriedades das Rochas de Reservatório"),
    ("Capítulo 4", "Integração de Dados, Método Volumétrico e EBM"),
    ("Apêndice",  "Perguntas & Exercícios por Capítulo"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}  —  {title}")
    run.font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 1 — SISTEMA PETROLÍFERO
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CAPÍTULO 1 — Sistema Petrolífero, Tipos de Reservatórios e Mecanismos de Produção",
            1, (0x1A, 0x52, 0x76))

add_heading("1.1 Sistema de Produção de Petróleo", 2)
add_para("O sistema de produção compreende o conjunto de equipamentos e processos desde o reservatório até à superfície: reservatório, poço, sistema de separação e tratamento, transporte e armazenamento.")

add_heading("1.2 Engenharia de Reservatórios", 2)
for item in [
    "Maximizar a eficiência de extracção dos hidrocarbonetos.",
    "Estimar volumes originais (OOIP/OGIP) e reservas.",
    "Definir estratégias de desenvolvimento e produção.",
    "Apoiar a tomada de decisão económica."
]:
    add_bullet(item)

add_heading("1.3 Tipos de Fluidos de Reservatório", 2)
styled_table(
    ["Tipo", "Características", "API / GOR"],
    [
        ("Óleo preto (Black oil)", "Baixo GOR, alta viscosidade", "< 40 °API"),
        ("Óleo volátil", "GOR moderado-alto, alta retracção", "35–45 °API"),
        ("Condensado retrógrado", "GOR > 5000 scf/STB, condensa na superfície", "> 45 °API"),
        ("Gás húmido", "Contém líquidos em superfície", "—"),
        ("Gás seco", "Praticamente só metano", "—"),
    ],
    header_bg="1A5276"
)

add_heading("1.4 Tipos de Reservatório", 2)
styled_table(
    ["Tipo", "Condição inicial", "Mecanismo principal"],
    [
        ("Undersaturated (subsaturado)", "p > pb  →  sem gás livre", "Expansão do óleo + compressibilidade"),
        ("Saturated (saturado)", "p = pb  →  início do gás livre", "Gás em solução (solution gas drive)"),
        ("Gas-cap (capa de gás)", "Gás livre no topo do reservatório", "Expansão da capa de gás"),
        ("Water drive", "Aquífero activo na base", "Influxo de água do aquífero"),
        ("Condensate", "T > Tc  →  gás no reservatório", "Condensação retrógrada"),
    ],
    header_bg="1A5276"
)

add_heading("1.5 Mecanismos de Produção Primária", 2)
mechs = [
    ("Expansão do óleo e água conata", "Domina em reservatórios subsaturados. Baixa recuperação (2–5 %)."),
    ("Gás em solução (solution gas drive)", "Gás dissolvido liberta-se com a queda de pressão. FR = 5–30 %."),
    ("Capa de gás (gas-cap drive)", "Expansão do gás livre empurra o óleo para baixo. FR = 20–40 %."),
    ("Influxo de água (water drive)", "Aquífero mantém pressão e desloca óleo. FR = 35–75 %."),
    ("Compactação da rocha", "Relevante em reservatórios de alta compressibilidade."),
    ("Gravidade (gravity drainage)", "Drenagem por gravidade em reservatórios de alta inclinação."),
]
for name, desc in mechs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"{name}: ")
    run.bold = True
    run.font.size = Pt(10.5)
    p.add_run(desc).font.size = Pt(10.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 2 — PROPRIEDADES DOS FLUIDOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CAPÍTULO 2 — Propriedades dos Fluidos de Reservatório", 1, (0x1E, 0x84, 0x49))

add_heading("2.1 Propriedades do Óleo", 2)

add_heading("2.1.1 Grau API e Densidade", 3)
add_box("°API = 141.5 / γo − 131.5      γo = ρóleo / ρágua@60°F (= 999.0 kg/m³)")
styled_table(
    ["Tipo de óleo", "°API", "γo"],
    [("Óleo extra-pesado", "< 10", "> 1.00"), ("Óleo pesado", "10–22", "0.92–1.00"),
     ("Óleo médio", "22–30", "0.88–0.92"), ("Óleo leve", "30–40", "0.83–0.88"),
     ("Óleo extra-leve", "> 40", "< 0.83")],
    header_bg="1E8449"
)

add_heading("2.1.2 Factor Volume de Formação do Óleo (Bo)", 3)
add_para("Bo [bbl/STB]: relação entre o volume do óleo no reservatório e o seu volume nas condições padrão (60°F, 1 atm). Aumenta com a pressão (acima de pb) e diminui abaixo de pb.")
add_box("Bo (acima de pb):  Bo = Boi · exp[−co(p − pi)]")

add_heading("2.1.3 Razão de Solubilidade do Gás (Rs)", 3)
add_para("Rs [scf/STB]: volume de gás dissolvido por STB de óleo nas condições de reservatório. Aumenta com a pressão até pb.")

add_heading("2.1.4 Factor Volume de Formação Bifásico (Bt)", 3)
add_box("Bt = Bo + (Rsi − Rs) · Bg")

add_heading("2.2 Propriedades do Gás", 2)
add_heading("2.2.1 Equação de Estado — Gás Real", 3)
add_box("pV = znRT      onde z = factor de compressibilidade (desvio do gás ideal)")

add_heading("2.2.2 Factor Volume de Formação do Gás (Bg)", 3)
add_box("Bg = 0.00504 · z · T / p    [bbl/scf],  T em °R,  p em psia")
add_box("Bg = 0.0283 · z · T / p     [ft³/scf]")

add_heading("2.2.3 Massa Molar Aparente e Densidade do Gás", 3)
add_box("Ma = 28.97 · γg       ρg = pMa / (zRT)     γg = ρg / ρar")

add_heading("2.2.4 Propriedades Pseudocríticas e Factor z (Standing-Katz)", 3)
add_box(
    "Tpc = Σ yi·Tci    ppc = Σ yi·pci\n"
    "Tpr = T / Tpc      ppr = p / ppc\n"
    "Correlações Standing:  γg < 0.75  →  Tpc ≈ 168+325γg−12.5γg²"
)

add_heading("2.2.5 Viscosidade do Gás", 3)
add_para("Aumenta com a pressão e a temperatura (ao contrário dos líquidos). Correlações: Lee-Gonzalez-Eakin, Carr-Kobayashi-Burrows.")

add_heading("2.3 Envelope de Fases", 2)
add_para("O envelope de fases define as condições (p, T) em que coexistem fases líquida e gasosa. O ponto crítico é onde as propriedades das duas fases se tornam idênticas. A curva de ponto de bolha separa o óleo sub-saturado da região bifásica; a curva de ponto de orvalho separa o gás saturado.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 3 — PROPRIEDADES DAS ROCHAS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CAPÍTULO 3 — Propriedades das Rochas de Reservatório", 1, (0x6C, 0x34, 0x83))

add_heading("3.1 Porosidade", 2)
add_box("φ = Vp / Vt      φ_total = (Vp_intergranular + Vp_fissu ras) / Vt")
styled_table(
    ["Classificação", "φ (%)"],
    [("Negligível", "< 5"), ("Pobre", "5–10"), ("Razoável", "10–15"),
     ("Boa", "15–20"), ("Muito boa", "> 20")],
    header_bg="6C3483"
)
add_para("Porosidade primária (intergranular): formada durante a deposição sedimentar. Porosidade secundária: criada por processos diagenéticos posteriores (dissolução, fracturação).")

add_heading("3.2 Compressibilidade", 2)
add_box(
    "cf = −(1/Vp)(dVp/dp)       [psi⁻¹]\n"
    "co = −(1/Vo)(dVo/dp)       cw = −(1/Vw)(dVw/dp)\n"
    "ct = So·co + Sw·cw + cf    [compressibilidade total]"
)

add_heading("3.3 Saturação de Fluidos", 2)
add_box("So + Sg + Sw = 1      Swi = saturação de água irredutível (connate)")
add_para("A saturação de água irredutível (Swi) é a mínima saturação de água que não é móvel. O óleo só flui acima da saturação crítica de óleo.")

add_heading("3.4 Permeabilidade (Lei de Darcy)", 2)
add_box(
    "FLUXO LINEAR:    q = kA(p1−p2)/(μL)\n"
    "FLUXO RADIAL:    q = 2πkh(pe−pw) / [μ ln(re/rw)]\n"
    "UNIDADES: q [cm³/s], k [Darcy], A [cm²], p [atm], μ [cP], L [cm]"
)
add_para("1 Darcy ≈ 9.869 × 10⁻¹³ m²   |   Reservatórios: 1–1000 mD típico")
styled_table(
    ["Tipo", "Descrição", "Símbolo"],
    [("Absoluta", "Único fluido saturando o meio", "k"),
     ("Efectiva", "Cada fase em presença das outras", "ko, kg, kw"),
     ("Relativa", "Efectiva / Absoluta (adimensional)", "kro, krg, krw")],
    header_bg="6C3483"
)

add_heading("3.5 Molhabilidade e Pressão Capilar", 2)
add_box(
    "cos θ = (σso − σsw) / σwo          [equilíbrio Young]\n"
    "pc = pnw − pw = 2σ cosθ / r       [tubo capilar]\n"
    "h = 2σ cosθ / (g·Δρ·r)            [altura capilar]\n"
    "J = pc·√(k/φ) / (σ cosθ)          [função J de Leverett]"
)
styled_table(
    ["Sistema", "θ (°)", "σ (dina/cm)", "σcosθ"],
    [("Ar – água (lab)", "0", "72", "72"),
     ("Óleo – água (lab)", "30", "48", "42"),
     ("Ar – mercúrio (lab)", "140", "480", "367"),
     ("Água – óleo (campo)", "30", "30", "26"),
     ("Água – gás (campo)", "0", "50", "50")],
    header_bg="6C3483"
)

add_heading("3.6 Embebição e Drenagem", 2)
add_para("Drenagem: deslocamento de fase molhante por fase não-molhante (saturação molhante diminui). Embebição: processo inverso. A histerese capilar resulta da diferença entre as curvas de drenagem e embebição.")
add_para("Pressão de deslocamento (pd): pressão mínima para iniciar a drenagem — corresponde ao poro de maior diâmetro.")

add_heading("3.7 Contacto de Hidrocarbonetos / Água", 2)
add_box(
    "FWL (Free Water Level): nível onde pc = 0  →  OWC ≠ FWL (OWC está acima por pc)\n"
    "ZFWL = (pw − po + ρo·g·Zo − ρw·g·Zw) / [(ρo − ρw)·g]"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 4 — MÉTODO VOLUMÉTRICO E EBM
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CAPÍTULO 4 — Integração de Dados, Método Volumétrico e EBM", 1, (0x92, 0x2B, 0x21))

add_heading("4.1 Definições Fundamentais", 2)
styled_table(
    ["Parâmetro", "Definição", "Fórmula"],
    [
        ("N (OOIP)", "Volume original de óleo", "equação volumétrica"),
        ("FR", "Factor de recuperação", "FR = NR / N"),
        ("NR", "Volume recuperável", "NR = N × FR"),
        ("NA / Np", "Produção acumulada", "medida"),
        ("fR", "Fracção recuperada", "fR = NA / N"),
        ("Reservas", "Volume ainda a produzir", "Reservas = NR − NA"),
    ],
    header_bg="922B21"
)

add_heading("4.2 Método Volumétrico", 2)
add_box(
    "OOIP:  N = 7758 · A · h · NTG · φ · (1 − Swi) / Boi     [STB]\n"
    "OGIP:  G = 43560 · A · h · NTG · φ · (1 − Swi) / Bgi   [scf]\n"
    "\n"
    "NTG = h_reservatório / h_total\n"
    "Cut-offs típicos: φ > 5–8 %;  Sw < 60–70 %;  k > 0.1–1 mD"
)
styled_table(
    ["Parâmetro", "Fonte de dados", "Unidade"],
    [("Área (A)", "Mapas sísmicos + petrofísica", "acres / m²"),
     ("Espessura (h)", "Registos de poço (GR, resistividade)", "ft / m"),
     ("NTG", "Petrofísica + cut-offs", "fracção"),
     ("Porosidade (φ)", "RCAL / registos de poço", "fracção"),
     ("Swi", "Registos de resistividade / Pc", "fracção"),
     ("Bo ou Bg", "Análise PVT", "bbl/STB ou bbl/scf")],
    header_bg="922B21"
)

add_heading("4.3 Classificação de Reservas (SPE/PRMS)", 2)
styled_table(
    ["Categoria", "Nível de certeza", "Probabilidade"],
    [("1P (Provadas — Proved)", "Alta", "P90  (≥ 90 %)"),
     ("2P (Provadas + Prováveis)", "Moderada", "P50  (≥ 50 %)"),
     ("3P (P+P+Possíveis)", "Baixa", "P10  (≥ 10 %)")],
    header_bg="922B21"
)

add_heading("4.4 Equação de Balanço de Materiais (EBM) Generalizada", 2)
add_box(
    "EBM: [Expansão dos fluidos + Injecção + Influxo] = [Produção acumulada]\n"
    "\n"
    "N = {Np[Bo + (Rp−Rs)Bg] − We − WpBw − GinjBginj − WinjBw} /\n"
    "    {(Bo−Boi) + (Rsi−Rs)Bg + mBoi(Bg/Bgi−1) + Boi(1+m)[(SwicW+cf)/(1−Swi)]Δp}"
)

add_heading("4.4.1 Linearização de Havlena-Odeh", 3)
add_box(
    "F = N·Eo + m·N·Eg + (1+m)·N·Ef,w + We\n"
    "\n"
    "Eo = Bo−Boi + (Rsi−Rs)·Bg                   [zona de óleo]\n"
    "Eg = Boi·(Bg/Bgi − 1)                        [capa de gás]\n"
    "Ef,w = Boi·(Swi·cw + cf)·Δp / (1−Swi)       [compressibilidades]\n"
    "F = Np·Bo + (Rp−Rs)·Bg + Wp·Bw − Winj·Bw − Ginj·Bginj"
)

add_heading("4.4.2 Casos Particulares da EBM", 3)
cases = [
    ("Gás em solução (p ≥ pb)",   "N = NpBo / {Boi·[(coSoi + cwSwi + cf)/Soi]·Δp}"),
    ("Gás em solução (p < pb)",   "F = N·Eo + N·Ef,w  →  F/N = Eo + Ef,w"),
    ("Capa de gás (m conhecido)", "F = N·Eo + m·N·Eg"),
    ("Water drive",               "F = N·Eo + We  →  F/Eo = N + We/Eo"),
    ("Water drive + capa de gás", "F = N·Eo + m·N·Eg + We"),
]
for name, formula in cases:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"{name}: ")
    run.bold = True
    run.font.size = Pt(10.5)
    add_box(formula)

add_heading("4.5 EBM para Reservatórios de Gás", 2)
add_box(
    "G = Gp·Bgi / (Bg − Bgi)         [sem aquífero, sem injecção]\n"
    "Linearizado (p/z plot):  p/z = pi/zi · (1 − Gp/G)\n"
    "Gráfico p/z × Gp → linha recta; intersecção com eixo Gp = G (OGIP)"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APÊNDICE — PERGUNTAS & EXERCÍCIOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("APÊNDICE — Perguntas & Exercícios por Capítulo", 1, (0x0D, 0x1B, 0x2A))
add_para("Compilação de todas as questões de consolidação, exercícios numéricos e tarefas das Aulas 0–12.")

# ── CAP 1 ──────────────────────────────────────────────────────────────────
add_heading("A.1 — Capítulo 1: Sistema Petrolífero", 2, (0x1A, 0x52, 0x76))

add_heading("A.1.1 Questões de Consolidação Aula 1 (SIM / NÃO)", 3)
styled_table(
    ["#", "Afirmação", "Resp."],
    [
        ("1","Sistema de produção: (A) Conjunto de equipamentos e processos desde o reservatório até superfície","SIM"),
        ("2","NÃO é função da Eng. Reservatórios: (C) Desenvolver tecnologias de perfuração","NÃO"),
        ("3","Reservatório undersaturated: (B) Pressão acima da saturação","SIM"),
        ("4","Objectivo principal: (A) Maximizar eficiência de extracção","SIM"),
        ("5","Gas-cap reservoir: (A) Gás livre sobre um óleo saturado","SIM"),
        ("6","Poço produtor: (A) Extracção de fluidos do reservatório","SIM"),
        ("7","Condensate reservoir: (A) Óleo que condensa a partir de gás com queda de pressão","SIM"),
        ("8","Eng. Reservatórios apoia decisões: (A) Permite estimar recuperação económica","SIM"),
        ("9","Diferença saturado vs subsaturado: saturado possui pressão ABAIXO de pb (gás livre)","SIM"),
        ("10","Manutenção de pressão: (A) Injecção de água ou gás","SIM"),
    ], header_bg="1A5276"
)

add_heading("A.1.2 Questões de Consolidação Aula 2 (SIM / NÃO)", 3)
styled_table(
    ["#", "Afirmação", "Resp."],
    [
        ("1","Óleo leve: (B) Baixa densidade e alta mobilidade","SIM"),
        ("2","Gás influencia recuperação: (A) Reduz pressão e facilita fluxo do óleo","SIM"),
        ("3","Water drive: (A) Manutenção de pressão por influxo de água do aquífero","SIM"),
        ("4","Viscosidade impacta: (A) A taxa de fluxo pelo reservatório","SIM"),
        ("5","Gas-cap drive: (A) Mantém pressão por expansão de gás acima do óleo","SIM"),
        ("6","Condensado: (A) Óleo líquido que condensa a partir do gás com queda de pressão","SIM"),
        ("7","Solution gas drive NÃO inclui produção contínua de óleo sem declínio de pressão","NÃO (inclui)"),
        ("8","Water influx: (A) Água do aquífero desloca óleo para os poços","SIM"),
        ("9","Undersaturated depende de: (A) Pressão inicial e mobilidade dos fluidos","SIM"),
        ("10","Gas-cap drive: (A) Gás livre no topo que expande e ajuda a manter pressão","SIM"),
    ], header_bg="1A5276"
)

add_heading("A.1.3 Tarefas / Questões Dissertativas", 3)
tasks_ch1 = [
    "Explique o conceito de sistema de produção e identifique os seus componentes principais.",
    "Descreva a importância da Engenharia de Reservatórios e as suas funções principais.",
    "Diferencie reservatórios 'undersaturated' e 'saturated' utilizando diagramas de fases.",
    "Como a história da Engenharia de Reservatórios influenciou os métodos modernos?",
    "Como o conhecimento do sistema petrolífero e dos tipos de reservatórios influencia o planeamento de desenvolvimento?",
    "Diferencie: Sistema de Produção, Sistema Petrolífero e cadeia produtiva do petróleo.",
    "Ilustre e explique o envelope de fases de um fluido de reservatório.",
    "Apresente as três teorias sobre a origem dos hidrocarbonetos, compare-as e indique a mais aceite.",
    "Descreva a área de actuação do Engenheiro de Reservatórios.",
    "Explique os mecanismos naturais de produção primária e compare expansão de gás vs. influxo de água.",
    "Construa uma tabela: tipo de reservatório × condições iniciais × comportamento com Δp × hidrocarbonetos produzidos.",
]
for i, t in enumerate(tasks_ch1, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(t).font.size = Pt(10.5)

# ── CAP 2 ──────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading("A.2 — Capítulo 2: Propriedades dos Fluidos", 2, (0x1E, 0x84, 0x49))

add_heading("A.2.1 Questões de Consolidação Aula 3 — Propriedades do Óleo", 3)
styled_table(
    ["#", "Tema", "Correct."],
    [
        ("1","Misturas homogéneas","A, B, E"),
        ("2","Soluções","A, B, D"),
        ("3","Misturas ideais","A, B, D"),
        ("4","Massa específica do óleo","A, B, D"),
        ("5","Peso específico relativo","A, B, D"),
        ("6","Densidade","A, B, D"),
        ("7","Grau API","A, B, C, E"),
        ("8","Compressibilidade isotérmica","A, B, D"),
        ("9","Factor Bo","A, B, D, E"),
        ("10","Viscosidade do óleo","A, B, D"),
    ], header_bg="1E8449"
)

add_heading("A.2.2 Questões de Consolidação Aula 4 — Propriedades do Gás", 3)
styled_table(
    ["#", "Tema", "Correct."],
    [
        ("1","Equação de estado dos gases","A, B, D"),
        ("2","Gases ideais","A, B, D, E"),
        ("3","Gases reais","A, B, C, E"),
        ("4","Factor Bg","A, B, D"),
        ("5","Viscosidade do gás","A, B, C, E"),
        ("6","Razão de solubilidade Rs","A, B, D"),
        ("7","Comportamento óleo-gás","A, B, C, E"),
        ("8","FVF bifásico Bt","A, B, D, E"),
        ("9","Equações de estado gases reais","A, B, C, E"),
        ("10","Propriedades de gases em reservatórios","A, B, D"),
    ], header_bg="1E8449"
)

add_heading("A.2.3 Exercícios Numéricos", 3)
exs_ch2 = [
    "Reservatório cilíndrico (D=2m, h=4m) com petróleo ρ=850 kg/m³: calcule a massa. Repita para gasolina ρ=740 kg/m³.",
    "ρ=740 kg/m³: determine o volume de 500 kg.",
    "400 kg em 1500 L: calcule ρ, peso específico e peso específico relativo.",
    "Óleo com ρ=53 lb/ft³: calcule γo e °API.",
    "Óleo de 42 °API: calcule γo e ρ (kg/m³).",
    "Gás a 510 ft³, 60°F, 12 psig: (a) volume a 110°F e mesma pressão; (b) nova pressão a volume constante e T=110°F.",
    "Preencher tanque de 100 ft³ a 40 psia / 90°F: (a) volume em condições padrão; (b) pressão ao arrefecer a 60°F.",
    "Poço de gás: γg=0.65, q=1.1 MMscf/dia, p=1500 psi, T=150°F. Calcule: (a) Ma; (b) ρg; (c) vazão mássica (lb/dia).",
    "Gás (CO₂=0.05, C₁=0.90, C₂=0.03, C₃=0.02): calcule factor z pelo método Standing-Katz (p=2000 psia, T=150°F).",
    "Calcule a massa do metano: V=5 ft³, p=700 psia, T=68°F.",
    "Gás ácido γg=0.75 com 5% CO₂ e 10% H₂S: calcule ρg a 3300 psia / 150°F (correcções Wichert-Aziz).",
    "p=2500 psia, T=180°F, z=0.85: (a) Bg [bbl/scf] e factor de expansão; (b) volume produzido para Vp=1×10⁹ ft³.",
    "Tabela e gráfico de Bg vs pressão para γg=0.64, T=81°F, de 3500 a 500 psia (passo 500 psia).",
]
for i, ex in enumerate(exs_ch2, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(ex).font.size = Pt(10.5)

add_heading("A.2.4 Tarefas / Questões Dissertativas", 3)
for t in [
    "Explique o papel das equações de estado na caracterização dos gases de reservatório.",
    "Compare o comportamento de gases ideais e reais em condições de reservatório.",
    "Defina razão de solubilidade (Rs) e explique a sua importância.",
    "Descreva o comportamento bifásico óleo-gás nos reservatórios.",
    "Fundamente as questões correctas e incorrectas das questões de consolidação.",
]:
    p = doc.add_paragraph(style='List Number')
    p.add_run(t).font.size = Pt(10.5)

# ── CAP 3 ──────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading("A.3 — Capítulo 3: Propriedades das Rochas", 2, (0x6C, 0x34, 0x83))

add_heading("A.3.1 Questões Aula 5 — Porosidade", 3)
styled_table(
    ["#", "Tema", "Correct."],
    [
        ("1","Rochas de reservatório","A, C, E"),
        ("2","Componentes da rocha","A, B, D"),
        ("3","Porosidade (definição)","A, B, D"),
        ("4","Porosidade primária","A, C, D"),
        ("5","Porosidade secundária","A, B, D, E"),
        ("6","Factores que influenciam porosidade","A, B, C, E"),
        ("7","Matriz rochosa","A, B, D, E"),
        ("8","Classificação da porosidade","A, B, D, E"),
        ("9","Alta porosidade","A, C, E"),
        ("10","Arenitos como reservatórios","A, B, D, E"),
    ], header_bg="6C3483"
)

add_heading("A.3.2 Questões Aula 6 — Compressibilidade, Saturação, Permeabilidade", 3)
styled_table(
    ["#", "Tema", "Correct."],
    [
        ("1","Compressibilidade da rocha","A, C, D"),
        ("2","Compressibilidade dos fluidos","A, C, E"),
        ("3","Saturação de fluidos","A, B, C, E"),
        ("4","Saturação de água","A, C, D"),
        ("5","Permeabilidade absoluta","A, C, D"),
        ("6","Permeabilidade relativa","A, B, D"),
        ("7","Factores que influenciam permeabilidade","A, B, C, E"),
        ("8","Compressibilidade total","A, B, D"),
        ("9","Saturação residual de óleo","A, B, D, E"),
        ("10","Relação entre propriedades","A, B, C, D"),
    ], header_bg="6C3483"
)

add_heading("A.3.3 Questões Aula 7 — Molhabilidade e Pressão Capilar", 3)
styled_table(
    ["#", "Tema", "Correct."],
    [
        ("1","Molhabilidade da rocha","A, C, D, E"),
        ("2","Pressão capilar","A, B, D, E"),
        ("3","Embebição e drenagem","A, C, E"),
        ("4","Função J de Leverett","A, B, D, E"),
        ("5","Molhabilidade e recuperação de óleo","D, E"),
        ("6","Pressão capilar e tamanho do poro","A, B, D"),
        ("7","Embebição/drenagem nas curvas capilares","A, C, D, E"),
        ("8","Função J e escalonamento","A, B, C, E"),
        ("9","Molhabilidade e eficiência de recuperação","A, C, E"),
        ("10","Processos capilares e fluxo multifásico","A, C, D, E"),
    ], header_bg="6C3483"
)

add_heading("A.3.4 Exercícios Numéricos", 3)
exs_ch3 = [
    "Amostra: Ws=130g, Wd=105g, ρóleo=840 kg/m³, Vt=180 cm³. Calcule a porosidade.",
    "Determine a porosidade idealizada para arranjos cúbico e romboédrico de esferas.",
    "Lei de Boyle: V1=100cc, V2=100cc, p1=15 psi, p2=60 psi, pf=39 psi. Calcule o volume do grão e a porosidade.",
    "Calcule a porosidade média de 4 amostras (média ponderada pela espessura).",
    "Wd=330g, Ws=360g, Wap(em água)=225g. Calcule a porosidade pelo método de Arquimedes.",
    "Complete a tabela de análise de rotina (RCAL): φ, k, Sw para cada amostra.",
    "Complete a tabela de análise especial (SCAL): curvas Pc, kr, molhabilidade.",
    "Vp=18 cm³, ΔVp=0.15 cm³, Δp=900 psig. Calcule cf.",
    "Vt=150 cm³, φ=18%, cf=7×10⁻⁶ psig⁻¹, grad. sobrecarga=0.95 psig/ft, prof.=8000 ft. Calcule φ a essa profundidade.",
    "Derive: 1 Darcy → m². Converta: (a) 3.5 mD em m²; (b) 7.3×10⁻¹² m² em Darcy.",
    "Permeabilidade: q=0.07 cm³/s, p_in=43 psig, L=3.5 in, D=1.48 in, μ=1 cP. Calcule k.",
    "Amostra: ρw=1 g/cm³, ρo=0.87 g/cm³, φ=13.6%, L=3 in, D=1.5 in, Ws=144.3g, Wd=133.2g. Calcule So e Sw.",
    "Derive a equação de fluxo radial permanente para fluido incompressível.",
    "Sistema (σwo=30, σso=80, σsw=65 dina/cm): (a) fluido que molha preferencialmente; (b) tensão de adesão; (c) ângulo θ.",
    "Poço com 3 tipos de rocha (A, B, C) e curvas Pc convertidas para altura: responda às alíneas a) a h).",
    "Medições Pc em laboratório, OWC=1530 m: (a) drenagem ou embebição?; (b) Swirr?; (c) converter para campo; (d) FWL?; (e) zona de transição.",
    "Testemunho φ=12%, k=50 mD, σ=40 dina/cm, θ=0°, Swirr=0.2; reservatório φ=16%, k=85 mD. Gere dados Pc representativos via função J.",
]
for i, ex in enumerate(exs_ch3, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(ex).font.size = Pt(10.5)

add_heading("A.3.5 Tarefas / Questões Dissertativas", 3)
for t in [
    "Explique rocha de reservatório e a sua importância.",
    "Diferencie porosidade total e efectiva.",
    "Descreva os efeitos da compactação na porosidade.",
    "Analise a influência dos processos diagenéticos nas propriedades da rocha.",
    "Explique compressibilidade e a sua importância. Diferencie cf da rocha e dos fluidos.",
    "Explique saturação de fluidos e analise a sua influência na permeabilidade relativa.",
    "Explique molhabilidade e descreva a pressão capilar e os factores que a influenciam.",
    "Explique embebição e drenagem e como influenciam estratégias de EOR.",
    "Fundamente as questões correctas e incorrectas das questões de consolidação.",
]:
    p = doc.add_paragraph(style='List Number')
    p.add_run(t).font.size = Pt(10.5)

# ── CAP 4 ──────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading("A.4 — Capítulo 4: Método Volumétrico e EBM", 2, (0x92, 0x2B, 0x21))

add_heading("A.4.1 Questões de Consolidação — Método Volumétrico", 3)
styled_table(
    ["#", "Tema", "Correct."],
    [
        ("1","Volume recuperável","B, C, D"),
        ("2","Sistemas em camadas","A, C, D"),
        ("3","Net to Gross (NTG)","A, B, C"),
        ("4","Cut-offs","A, B, C, E"),
        ("5","Integração de dados","A, C, D"),
        ("6","Incertezas no cálculo volumétrico","A, B, D, E"),
        ("7","Parâmetros para volume original","A, B, C, D"),
        ("8","Parâmetros para caracterizar o fluxo","A, B, C, E"),
        ("9","Relação cut-offs e reservas","A, B, D"),
        ("10","Sistemas em camadas e fluxo","A, C, D"),
    ], header_bg="922B21"
)

add_heading("A.4.2 Questões de Consolidação — EBM (Aulas 11-12)", 3)
styled_table(
    ["#", "Tema", "Correct."],
    [
        ("1","Base da EBM","A, C, E"),
        ("2","Utilizações da EBM","A, B, D"),
        ("3","Equação generalizada","A, B, C, D"),
        ("4","Reservatório como tanque","A, B, C, E"),
        ("5","Limitações da EBM","A, B, E"),
        ("6","EBM — identificação de mecanismos","A, B, D, E"),
        ("7","Pressão média do reservatório","A, B, D"),
        ("8","Dados necessários para EBM","A, B, D, E"),
        ("9","Expansão dos fluidos","A, C"),
        ("10","Aplicabilidade da EBM","A, B, D"),
    ], header_bg="922B21"
)

add_heading("A.4.3 Exercícios Numéricos — Método Volumétrico", 3)
exs_vol = [
    "Reservatório homogéneo com Bo=1.31 bbl/STB, Swi=20%: (a) calcule N pelo método volumétrico; (b) calcule o NTG.",
    "Investigue os 20 países com maiores reservas de óleo e gás no mundo (tabela comparativa).",
    "Reservatório de gás: A=3000 acres, h=30 ft, φ=15%, Swi=20%, T=150°F, pi=2600 psi, γg=0.75, pab=400 psi. Gráficos: (a) Gp vs Bg; (b) Gp vs z; (c) Gp vs p.",
    "N=3.2×10⁶ Sm³; VR=7.36×10⁵ Sm³; NA(3anos)=4×10⁵ Sm³. Calcule: (a) FR; (b) fR; (c) Reservas.",
]
for ex in exs_vol:
    p = doc.add_paragraph(style='List Number')
    p.add_run(ex).font.size = Pt(10.5)

add_heading("A.4.4 Exercícios Numéricos — EBM", 3)
exs_ebm = [
    "Campo Big Butte (mecanismo combinado): p=2500 psia; Vóleo=100000 ft³; Vgás=20000 ft³; dados PVT em tabela (pi=3000 psia, Boi=1.35, Gp=5.5 MMMscf, Wp=0.2 MMbbl, Bg=0.0015). Calcule N.",
    "Vp=75 MMft³, sem capa de gás e sem aquífero; Rs=0.42 scf/STB/psi; pi=3500 psia; pb=2400 psia; Boi=1.33; z=0.95 a 1500 psia; Np=1.0 MMSTB; GOR=2800 scf/STB. Calcule: (a) N; (b) G inicial; (c) GOR dissolvida inicial; (d) gás remanescente; (e) gás livre; (f) volume do reservatório de gás livre; (g) GOR total a 1500 psia.",
    "pi=3685 psia; cf=4.95×10⁻⁶ psi⁻¹; cw=3.62×10⁻⁶ psi⁻¹; Swi=24%; Bw=1.0; pb=1500 psi. Dados de produção e PVT em tabela (13 pressões). Use EBM para calcular N.",
    "pb=234.18 kgf/cm²; sem água; dados PVT (6 pressões). Determine: (a) volume original de óleo; (b) volume original de gás capa.",
    "pb=234.18 kgf/cm²; cf=122.3×10⁻⁶ (kgf/cm²)⁻¹; cw=42.7×10⁻⁶; Swi=0.20; pab=63.3 kgf/cm². Determine: (a) fR até pb; (b) expressão para fR como função de Rps até ao abandono; (c) saturação de gás livre para recuperação final de 12%.",
    "pi=2500 psia; GOR acumulada=954 scf/STB; sem aquífero nem capa de gás. Dados PVT (4 pressões). Calcule N.",
    "N=650 MMSTB; pi=7150 psia; pb=4500 psia; Boi=1.743; Bob=1.850; cw=3.0×10⁻⁶/psi; cf=3.3×10⁻⁶/psi; Swi=0.43; Npb=43.473 MMSTB. Calcule: (a) N inicial e saturação de gás após 72 meses; (b) variação do volume de óleo sem cw e cf.",
]
for ex in exs_ebm:
    p = doc.add_paragraph(style='List Number')
    p.add_run(ex).font.size = Pt(10.5)

add_heading("A.4.5 Tarefas / Questões Dissertativas", 3)
for t in [
    "Discuta o conceito de volume recuperável e a sua importância.",
    "Analise o comportamento de reservatórios em sistemas estratificados.",
    "Explique NTG e o papel dos cut-offs na avaliação de reservatórios.",
    "Discuta a importância da integração de dados na caracterização do fluxo.",
    "Explique o princípio fundamental da EBM.",
    "Descreva os principais termos da equação generalizada de balanço de materiais.",
    "Discuta a pressão média do reservatório na aplicação da EBM.",
    "Explique a linearização da EBM (método de Havlena-Odeh).",
    "Explique a gestão da injecção de água em reservatórios de óleo.",
    "Explique o comportamento do reservatório acima e abaixo da pressão de bolha.",
    "Compare EBM, método volumétrico e simulação numérica.",
    "Fundamente as questões correctas e incorrectas das questões de consolidação.",
]:
    p = doc.add_paragraph(style='List Number')
    p.add_run(t).font.size = Pt(10.5)

# ── FÓRMULAS ESSENCIAIS ─────────────────────────────────────────────────────
doc.add_page_break()
add_heading("Fórmulas Essenciais para Exame", 1, (0x0D, 0x1B, 0x2A))
styled_table(
    ["Área", "Fórmula / Expressão"],
    [
        ("°API / Densidade óleo", "°API = 141.5/γo − 131.5   |   γo = ρóleo/ρágua@60°F"),
        ("Gás real", "pV = znRT   |   z obtido por Standing-Katz (Tpr, ppr)"),
        ("Bg", "Bg = 0.00504·zT/p  [bbl/scf]   |   T em °R, p em psia"),
        ("FVF bifásico", "Bt = Bo + (Rsi − Rs)·Bg"),
        ("Porosidade", "φ = Vp/Vt   |   Lei de Boyle: (p1V1 + p2V2) = (p1+p2)Vt + ... "),
        ("Saturação", "So + Sg + Sw = 1"),
        ("Darcy linear", "q = kA(p1−p2)/(μL)"),
        ("Darcy radial", "q = 2πkh(pe−pw)/[μ ln(re/rw)]"),
        ("Pressão capilar", "pc = pnw − pw = 2σcosθ/r   |   h = 2σcosθ/(gΔρr)"),
        ("Função J", "J = pc√(k/φ)/(σcosθ)"),
        ("OOIP volumétrico", "N = 7758·A·h·NTG·φ·(1−Swi)/Boi  [STB]"),
        ("OGIP volumétrico", "G = 43560·A·h·NTG·φ·(1−Swi)/Bgi  [scf]"),
        ("EBM — Havlena-Odeh", "F = N·Eo + m·N·Eg + (1+m)·N·Ef,w + We"),
        ("Eo", "Eo = Bo−Boi + (Rsi−Rs)·Bg"),
        ("Eg", "Eg = Boi·(Bg/Bgi − 1)"),
        ("Ef,w", "Ef,w = Boi·(Swi·cw+cf)·Δp/(1−Swi)"),
        ("p/z plot (gás)", "p/z = (pi/zi)·(1 − Gp/G)   →   intersecção = G"),
        ("Reservas", "Reservas = N·FR − NA"),
    ],
    header_bg="1A1A2E", header_fg="F0F0F0"
)

# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"✔ DOCX saved: {OUT}")
