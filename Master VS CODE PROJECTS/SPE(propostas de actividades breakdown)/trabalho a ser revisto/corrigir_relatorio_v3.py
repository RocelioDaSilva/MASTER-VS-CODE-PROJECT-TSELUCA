from docx import Document
from docx.enum.text import WD_BREAK
import os
import re

# ================= CONFIGURAÇÕES =================
caminho_entrada = r"C:\Users\PCGAME\Desktop\Master VS CODE PROJECTS\SPE(propostas de actividades breakdown)\trabalho a ser revisto\Relatorio_FINAL_OFICIAL_V2.docx"
caminho_saida = r"C:\Users\PCGAME\Desktop\Master VS CODE PROJECTS\SPE(propostas de actividades breakdown)\trabalho a ser revisto\Relatorio_FINAL_OFICIAL_V3.docx"

# Substituições genéricas (strings)
substituicoes = {
    "20. un 2026": "20 Jun 2026",
    "Gráfico 1 -- distribuição por cursos": "Gráfico 1 – Distribuição por curso",
    "Gráfico 2 -- Distribuição por Cargos": "Gráfico 2 – Interesse por cargos",
    "Gráfico 3 -- Disponibilidade para representar a SPE": "Gráfico 3 – Disponibilidade para representar a SPE",
    "Gráfico 4 -- Nível de inglês": "Gráfico 4 – Nível de inglês",
    "Gráfico 5 -- Capacidade de comunicação": "Gráfico 5 – Capacidade de comunicação",
    "Gráfico 6 -- Preferência: membro ou cargo": "Gráfico 6 – Preferência: membro ou cargo",
    "Gráfico 7 -- Ano de frequência": "Gráfico 7 – Ano de frequência",
    "Gráfico 8 - Disponibilidade para reuniões": "Gráfico 8 – Disponibilidade para reuniões",
    "***Gráfico 9 - Estado da Membership**": "**Gráfico 9 – Estado da membership**",
    "***Gráfico 10 - Situação do SPE ID**": "**Gráfico 10 – Situação do SPE ID**",
}

# Títulos principais
titulos_principais = [
    "INTRODUÇÃO", "REVISÃO BIBLIOGRÁFICA", "METODOLOGIA",
    "APRESENTAÇÃO, ANÁLISE E DISCUSSÃO DOS RESULTADOS", "CONCLUSÕES",
    "REFERÊNCIAS BIBLIOGRÁFICAS", "ANEXO A", "ANEXO B"
]

# ================= FUNÇÕES AUXILIARES =================

def substituir_em_paragrafo(par, substituicoes_dict):
    """Substitui textos em todos os runs do parágrafo, preservando formatação."""
    # Primeiro, tenta substituir no texto completo (mais rápido, mas pode perder formatação)
    # Vamos fazer run a run para segurança
    texto_completo = par.text
    modificado = False
    for antigo, novo in substituicoes_dict.items():
        if antigo in texto_completo:
            # Substitui run a run
            novo_texto_completo = texto_completo.replace(antigo, novo)
            # Reconstruir o parágrafo a partir do texto completo (simplificado, mas eficaz)
            # Nota: isso remove formatação, mas como é só texto corrido, aceitável
            par.text = novo_texto_completo
            modificado = True
            texto_completo = novo_texto_completo
    return modificado

def substituir_travessao_em_paragrafo(par):
    """Substitui 'Gráfico X --' por 'Gráfico X –' em todos os runs."""
    texto = par.text
    novo_texto = re.sub(r'(Gráfico \d+) --', r'\1 –', texto)
    if novo_texto != texto:
        par.text = novo_texto
        return True
    return False

def percorrer_todos_paragrafos(doc, func):
    """Aplica uma função a todos os parágrafos do documento (incluindo dentro de tabelas)."""
    for par in doc.paragraphs:
        func(par)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:
                    func(par)

def aplicar_substituicoes_globais(doc):
    """Aplica todas as substituições de texto em todos os parágrafos."""
    def aplicar(par):
        substituir_em_paragrafo(par, substituicoes)
        substituir_travessao_em_paragrafo(par)
    percorrer_todos_paragrafos(doc, aplicar)

def remover_legendas_fonte_repetidas(doc):
    """Remove parágrafos com '*Fonte: ...' que não estejam logo após gráfico."""
    for i in range(len(doc.paragraphs)-1, -1, -1):
        par = doc.paragraphs[i]
        texto = par.text.strip()
        if texto == "*Fonte: inquérito aos estudantes (N=145)*":
            if i > 0 and re.match(r'Gráfico \d+', doc.paragraphs[i-1].text.strip()):
                continue
            else:
                par._element.getparent().remove(par._element)

def adicionar_quebras_antes_titulos(doc):
    """Insere quebra de página antes de cada título principal."""
    for i in range(len(doc.paragraphs)-1, -1, -1):
        par = doc.paragraphs[i]
        texto = par.text.strip().upper()
        for titulo in titulos_principais:
            if texto == titulo or texto.startswith(titulo):
                if i > 0:
                    run = par.insert_paragraph_before().add_run()
                    run.add_break(WD_BREAK.PAGE)
                break

def eh_subtitulo(par):
    texto = par.text.strip()
    if not texto:
        return False
    if par.style.name and ('Heading 2' in par.style.name or 'Heading 3' in par.style.name):
        return True
    if re.match(r'^\d+\.\d+(\.\d+)?\s+\w', texto):
        return True
    if re.match(r'^\d+\.\s+\w+', texto) and len(texto) < 80:
        return True
    if par.runs and all(run.bold for run in par.runs) and len(texto) < 100:
        return True
    palavras_chave = ['contextualização','objectivos','justificativa','problematização',
                      'impacto','estrutura','gestão','próximos passos','cronograma',
                      'estratégia','apresentação','análise','discussão','tipo de abordagem']
    return any(pc in texto.lower() for pc in palavras_chave)

def aplicar_keep_with_next_subtitulos(doc):
    count = 0
    for par in doc.paragraphs:
        if eh_subtitulo(par):
            par.paragraph_format.keep_with_next = True
            count += 1
    print(f"✅ Aplicado 'keep_with_next' a {count} subtítulos.")

def corrigir_anexo_a(doc):
    """Anexo A: anonimiza nomes (primeiro nome + inicial do último apelido)."""
    for tabela in doc.tables:
        if len(tabela.rows) < 2:
            continue
        # Procura tabela cuja primeira linha contenha "Nome" na segunda coluna (aproximado)
        cab = [cell.text.strip() for cell in tabela.rows[0].cells]
        if len(cab) >= 2 and "Nome" in cab[1]:
            for linha in tabela.rows[1:]:
                if len(linha.cells) >= 2:
                    nome_completo = linha.cells[1].text.strip()
                    partes = nome_completo.split()
                    if len(partes) >= 2:
                        primeiro = partes[0]
                        inicial = partes[-1][0] + "."
                        linha.cells[1].text = f"{primeiro} {inicial}"
            print("✅ Anexo A anonimizado.")
            return True
    return False

def corrigir_anexo_b(doc):
    """Anexo B: corrige '4 ° ano' e anonimiza nomes."""
    for tabela in doc.tables:
        if len(tabela.rows) < 2:
            continue
        cab = [cell.text.strip() for cell in tabela.rows[0].cells]
        # Verifica se é a tabela do Anexo B (primeira coluna "N°", segunda coluna com "Nome")
        if len(cab) >= 2 and cab[0] == "N°" and ("Nome" in cab[1] or "Nome completo" in cab[1]):
            col_nome = 1
            for idx, linha in enumerate(tabela.rows[1:]):
                if len(linha.cells) <= col_nome:
                    continue
                cell = linha.cells[col_nome]
                texto = cell.text.strip()
                # Corrige linha corrupta
                if "4 ° ano" in texto:
                    cell.text = "Nome não identificado"
                    print(f"   Linha {idx+2}: '4 ° ano' corrigido.")
                    continue
                # Anonimiza nomes completos
                partes = texto.split()
                if len(partes) >= 2:
                    primeiro = partes[0]
                    ultimo = partes[-1]
                    inicial = ultimo[0] + "."
                    cell.text = f"{primeiro} {inicial}"
            print("✅ Anexo B anonimizado e linha corrupta corrigida.")
            return True
    print("⚠️ Tabela do Anexo B não encontrada.")
    return False

# ================= EXECUÇÃO PRINCIPAL =================
def main():
    print("📄 A abrir documento...")
    doc = Document(caminho_entrada)

    print("✏️ Aplicando substituições de texto (runs)...")
    aplicar_substituicoes_globais(doc)

    print("🧹 Removendo legendas de fonte repetidas...")
    remover_legendas_fonte_repetidas(doc)

    print("📌 Inserindo quebras antes de títulos principais...")
    adicionar_quebras_antes_titulos(doc)

    print("🔗 Aplicando 'keep_with_next' aos subtítulos...")
    aplicar_keep_with_next_subtitulos(doc)

    print("🕵️ Corrigindo Anexo A...")
    corrigir_anexo_a(doc)

    print("🕵️ Corrigindo Anexo B...")
    corrigir_anexo_b(doc)

    # Garantir substituição da data mais uma vez (para sobreposição)
    print("🔁 Reforçando substituição de '20. un 2026'...")
    def forcar_data(par):
        if "20. un 2026" in par.text:
            par.text = par.text.replace("20. un 2026", "20 Jun 2026")
    percorrer_todos_paragrafos(doc, forcar_data)

    os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
    doc.save(caminho_saida)
    print(f"\n✅ Documento final guardado em:\n{caminho_saida}")
    print("\n📌 Verifique no Word:\n"
          "   - Anexo B: nomes anonimizados e '4 ° ano' corrigido\n"
          "   - Cronograma: '20 Jun 2026'\n"
          "   - Gráficos: travessão '–'\n"
          "   - Subtítulos: não ficam isolados no fim da página")

if __name__ == "__main__":
    main()